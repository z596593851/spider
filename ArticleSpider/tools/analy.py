import re
from docx import Document
from pdfminer.converter import PDFPageAggregator
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.layout import LTTextBoxHorizontal, LAParams
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter

d_pattern = re.compile('一、|二、|三、|四、|五、|六、|七、|八、|九、|十、|'
                       '$（一）|$（二）|$（三）|$（四）|$（五）|$（六）|$（七）|$（八）|$（九）|$（十）|$（十一）|$（十二）|$（十三）|$（十四）|$（十五）')
x_pattern=re.compile('1、|2、|3、|4、|5、|6、|7、|8、|9、|0、|'
                     '1[.．]\D|2[.．]\D|3[.．]\D|4[.．]\D|5[.．]\D|6[.．]\D|7[.．]\D|8[.．]\D|9[.．]\D|0[.．]\D')


def analy_doc(path):
    suffix=path.split(".")[-1]
    # print(suffix)
    # page_count=0
    # word_count=0
    if suffix == 'pdf':
        return parser_pdf_file(path)
        # page_count, word_count=parser_pdf_file(path)
        # return page_count, word_count
    if suffix == 'docx':
        return parser_word_file(path)
        # page_count, word_count=parser_word_file(path)
        # return page_count, word_count




def parser_word_file(word_file_path):
    d_count=0
    x_count=0
    word_count=0

    document = Document(word_file_path)
    for paragraph in document.paragraphs:
        result=paragraph.text.strip()
        d_match = d_pattern.findall(result)
        x_match = x_pattern.findall(result)
        word_count+=len(result)
        print(result)
        print("======")

        if d_match:
            d_count+=1
        if x_match:
            x_count+=1
    if d_count == 0:
        return x_count,word_count
    else:
        return d_count, word_count

def parser_pdf_file(pdf_file_path):
    d_count = 0
    x_count = 0
    word_count=0
    read_pdf = open(pdf_file_path, 'rb')  # 打开PDF文件。
    parser_pdf = PDFParser(read_pdf)  # 用文件对象创建一个PDF文档分析器。
    pdf_document = PDFDocument(parser_pdf)  # 创建一个PDF文档。

    parser_pdf.set_document(pdf_document)
    pdf_document.set_parser(parser_pdf)  # 连接分析器 与文档对象。
    pdf_document.initialize()  # 如果没有密码，就创建一个空的字符串。

    if not pdf_document.is_extractable:  # 检测文档是否提供txt转换，不提供就忽略。
        raise PDFTextExtractionNotAllowed
    else:
        pdf_manager = PDFResourceManager()  # 创建PDF资源管理器 来管理共享资源。
        pdf_laparams = LAParams()  # 创建一个PDF参数分析器。
        pdf_device = PDFPageAggregator(pdf_manager, laparams=pdf_laparams)  # 创建一个聚合器
        pdf_interpreter = PDFPageInterpreter(pdf_manager, pdf_device)  # 创建一个PDF页面解释器对象
        # 循环遍历列表，每次处理一页的内容，pdf_document.get_pages()获取page列表
        for each_page in pdf_document.get_pages():
            pdf_interpreter.process_page(each_page)  # 使用页面解释器来读取
            layout = pdf_device.get_result()  # 这里layout是一个LTPage对象 里面存放着这个page解析出的各种对象 一般包括LTTexBox,LTFigure,LTImage,
            # LTTexBoxHorizontal等等 想要获取文本就获得对象的text属性。
            # print(layout)
            for each_info in layout:
                if isinstance(each_info, LTTextBoxHorizontal):
                    result = each_info.get_text().strip()
                    d_match = d_pattern.findall(result)
                    x_match = x_pattern.findall(result)

                    word_count += len(result)

                    if d_match:
                        d_count += 1
                    if x_match:
                        x_count += 1
                    print(result)
                    print("======")
        if d_count == 0:
            return x_count,word_count
        else:
            return d_count,word_count

if __name__ == '__main__':
    page_count, word_count=analy_doc(r'C:\Users\huxiaoming02\Desktop\11e8456a6dfafba2b3e98f74f91929ff1055f09c.docx')
    print(page_count,word_count)